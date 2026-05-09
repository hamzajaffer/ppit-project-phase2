"""
DOCX Generator Module
Creates formatted Word documents from analyzed text blocks.
Applies heading styles, bold/italic formatting, alignment, and paragraph structure.

Phase 2: Added bullet list support and agent metadata appendix.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict, Any
import io


def generate_docx(
    text_blocks: List[Dict[str, Any]],
    title: str = "Converted Document",
    agent_metadata: Dict[str, Any] = None,
) -> io.BytesIO:
    """Generate a formatted .docx file from analyzed text blocks."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Add document title
    doc.add_heading(title, level=0)

    current_paragraph = None

    for block in text_blocks:
        text = block['text'].strip()
        if not text:
            continue

        is_heading = block.get('is_heading', False)
        is_bold = block.get('is_bold', False)
        is_bullet = block.get('is_bullet', False)
        alignment = block.get('alignment', 'left')
        new_paragraph = block.get('new_paragraph', True)
        font_size_cat = block.get('font_size_category', 'normal')

        if is_heading:
            heading_level = 1 if font_size_cat == 'large' else 2
            para = doc.add_heading(text, level=heading_level)
            set_paragraph_alignment(para, alignment)
            current_paragraph = None

        elif is_bullet:
            # Phase 2: Bullet list support
            # Strip the bullet character
            clean_text = text.lstrip('•●○■□▪▫‣⁃-–— ')
            # Remove numbered prefix like "1. " or "a. "
            if clean_text and clean_text[0].isdigit():
                dot_pos = clean_text.find('.')
                if 0 < dot_pos <= 3:
                    clean_text = clean_text[dot_pos+1:].strip()
            elif len(clean_text) > 2 and clean_text[0].isalpha() and clean_text[1] == '.':
                clean_text = clean_text[2:].strip()

            para = doc.add_paragraph(clean_text, style='List Bullet')
            run = para.runs[0] if para.runs else para.add_run(clean_text)
            apply_run_formatting(run, is_bold, font_size_cat)
            current_paragraph = None

        elif new_paragraph or current_paragraph is None:
            para = doc.add_paragraph()
            set_paragraph_alignment(para, alignment)
            run = para.add_run(text)
            apply_run_formatting(run, is_bold, font_size_cat)
            current_paragraph = para

        else:
            run = current_paragraph.add_run(' ' + text)
            apply_run_formatting(run, is_bold, font_size_cat)

    # If document is empty, add a note
    if len(doc.paragraphs) <= 1:
        doc.add_paragraph(
            "No text could be extracted from the image. "
            "This may be due to image quality or handwriting complexity."
        )

    # Phase 2: Add agent metadata appendix if available
    if agent_metadata:
        _add_agent_appendix(doc, agent_metadata)

    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _add_agent_appendix(doc: Document, metadata: Dict[str, Any]):
    """Add an agent processing report as an appendix to the document."""
    doc.add_page_break()
    doc.add_heading('Agent Processing Report', level=1)

    # Analysis summary
    analysis = metadata.get('analysis', {})
    if analysis:
        doc.add_heading('Image Analysis', level=2)
        for key, value in analysis.items():
            para = doc.add_paragraph()
            run = para.add_run(f"{key}: ")
            run.bold = True
            para.add_run(str(value))

    # Strategy used
    strategy = metadata.get('strategy', {})
    if strategy:
        doc.add_heading('Strategy Used', level=2)
        para = doc.add_paragraph()
        run = para.add_run(f"{strategy.get('name', 'N/A')}")
        run.bold = True
        para.add_run(f" — {strategy.get('description', '')}")

    # Decisions
    decisions = metadata.get('decisions', [])
    if decisions:
        doc.add_heading('Agent Decisions', level=2)
        for d in decisions:
            para = doc.add_paragraph(style='List Bullet')
            run = para.add_run(d.get('message', ''))
            run.bold = True
            if d.get('reasoning'):
                para.add_run(f"\n  Reasoning: {d['reasoning']}")

    # Quality info
    doc.add_heading('Quality Info', level=2)
    para = doc.add_paragraph()
    para.add_run(f"Retries: {metadata.get('retries', 0)}")
    para = doc.add_paragraph()
    passed = metadata.get('quality_passed', False)
    para.add_run(f"Quality Gate: {'PASSED' if passed else 'PARTIAL'}")


def apply_run_formatting(run, is_bold: bool, font_size_cat: str):
    """Apply formatting to a text run."""
    if is_bold:
        run.bold = True
        run.font.color.rgb = RGBColor(180, 0, 0)

    if font_size_cat == 'large':
        run.font.size = Pt(14)
    elif font_size_cat == 'small':
        run.font.size = Pt(9)
    else:
        run.font.size = Pt(11)

    run.font.name = 'Calibri'


def set_paragraph_alignment(paragraph, alignment: str):
    """Set paragraph alignment."""
    if alignment == 'center':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == 'right':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def generate_plain_docx(text: str, title: str = "Converted Document") -> io.BytesIO:
    """Generate a simple .docx with plain text (fallback mode)."""
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    doc.add_heading(title, level=0)

    paragraphs = text.split('\n\n')
    for para_text in paragraphs:
        lines = para_text.strip().split('\n')
        for line in lines:
            line = line.strip()
            if line:
                if line.startswith('#'):
                    clean = line.lstrip('#').strip()
                    hash_count = line.count('#', 0, line.index(' ') if ' ' in line else len(line))
                    level = max(1, min(hash_count, 4))
                    doc.add_heading(clean, level=level)
                else:
                    doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
